# ingestion/loader.py
from pathlib import Path
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_core.documents import Document

# Supported file extensions
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


def load_documents(file_path: str):
    """Load documents from file. Supports TXT, PDF, and DOCX."""
    path = Path(file_path)
    ext = path.suffix.lower()
    
    if ext == ".txt":
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()
    
    elif ext == ".pdf":
        loader = PyPDFLoader(file_path)
        return loader.load()
    
    elif ext == ".docx":
        return _load_docx(file_path)
    
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {SUPPORTED_EXTENSIONS}")


def _load_docx(file_path: str):
    """Load DOCX file using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required for DOCX files. Install with: pip install python-docx")
    
    doc = DocxDocument(file_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return [Document(page_content=text, metadata={"source": file_path})]


def is_supported_file(filename: str) -> bool:
    """Check if file extension is supported."""
    ext = Path(filename).suffix.lower()
    return ext in SUPPORTED_EXTENSIONS
